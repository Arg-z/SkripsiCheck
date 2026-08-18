"""Safe, local-only text extraction for supported document formats."""

from __future__ import annotations

import zipfile
from pathlib import Path

import pymupdf
from docx import Document

from app.config import SETTINGS

SUPPORTED_EXTENSIONS = frozenset({".pdf", ".docx", ".txt"})


class DocumentExtractionError(ValueError):
    """Raised when a document is invalid, unsupported, or unreadable."""


def _validate_document(path: Path, max_size_mb: int) -> Path:
    candidate = path.expanduser().resolve(strict=True)
    if not candidate.is_file():
        raise DocumentExtractionError(f"Not a regular file: {candidate.name}")

    suffix = candidate.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        allowed = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise DocumentExtractionError(f"Unsupported file type. Allowed: {allowed}")

    if candidate.stat().st_size > max_size_mb * 1024 * 1024:
        raise DocumentExtractionError(f"File exceeds the {max_size_mb} MB limit.")

    with candidate.open("rb") as document_file:
        signature = document_file.read(8)
    if suffix == ".pdf" and not signature.startswith(b"%PDF-"):
        raise DocumentExtractionError("The file extension is PDF but its content is not PDF.")
    if suffix == ".docx":
        if not signature.startswith(b"PK") or not zipfile.is_zipfile(candidate):
            raise DocumentExtractionError("The file extension is DOCX but its content is not DOCX.")
        with zipfile.ZipFile(candidate) as archive:
            if "word/document.xml" not in archive.namelist():
                raise DocumentExtractionError("The DOCX package has no Word document content.")
    if suffix == ".txt" and b"\x00" in signature:
        raise DocumentExtractionError("Binary content is not accepted as a text document.")
    return candidate


def _extract_pdf(path: Path) -> str:
    with pymupdf.open(path) as document:
        if document.needs_pass:
            raise DocumentExtractionError("Password-protected PDFs are not supported.")
        # Form feed preserves page boundaries for header/footer cleaning.
        return "\f".join(page.get_text("text") for page in document)


def _extract_docx(path: Path) -> str:
    document = Document(path)
    blocks: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))
    return "\n\n".join(blocks)


def _extract_txt(path: Path) -> str:
    raw = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-16", "cp1252"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentExtractionError("Text encoding is not supported.")


def extract_text(path: str | Path, max_size_mb: int | None = None) -> str:
    """Extract text without executing or uploading document contents."""

    limit = max_size_mb if max_size_mb is not None else SETTINGS.max_upload_mb
    try:
        validated = _validate_document(Path(path), limit)
        extractor = {
            ".pdf": _extract_pdf,
            ".docx": _extract_docx,
            ".txt": _extract_txt,
        }[validated.suffix.lower()]
        return extractor(validated)
    except DocumentExtractionError:
        raise
    except (OSError, RuntimeError, ValueError, zipfile.BadZipFile) as exc:
        raise DocumentExtractionError(f"Could not extract {Path(path).name}: {exc}") from exc
